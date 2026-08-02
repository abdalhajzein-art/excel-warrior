import sys
import os
import json
import pandas as pd

# استيراد آمن للمكتبات الأخرى
try:
    import docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

def extract_metadata(file_path):
    if not os.path.exists(file_path):
        return {"error": "File not found"}

    ext = os.path.splitext(file_path)[1].lower()
    size_bytes = os.path.getsize(file_path)

    meta = {
        "file_name": os.path.basename(file_path),
        "extension": ext,
        "size_kb": round(size_bytes / 1024, 2),
        "type": "unknown",
        "details": {}
    }

    try:
        # 1. معالجة ملفات الإكسل (Excel)
        if ext in ['.xlsx', '.xls', '.xlsm']:
            meta["type"] = "excel"
            excel_file = pd.ExcelFile(file_path)
            sheets_meta = {}
            for sheet_name in excel_file.sheet_names:
                df_sample = pd.read_excel(file_path, sheet_name=sheet_name, nrows=5)
                full_shape = pd.read_excel(file_path, sheet_name=sheet_name).shape
                
                sheets_meta[sheet_name] = {
                    "total_rows": full_shape[0],
                    "total_columns": full_shape[1],
                    "columns": list(df_sample.columns),
                    "dtypes": {col: str(dtype) for col, dtype in df_sample.dtypes.items()},
                    "sample_rows": df_sample.fillna("").to_dict(orient="records")
                }
            meta["details"]["sheets"] = sheets_meta

        # 2. معالجة ملفات (CSV)
        elif ext == '.csv':
            meta["type"] = "csv"
            df_sample = pd.read_csv(file_path, nrows=5)
            full_shape = pd.read_csv(file_path).shape
            meta["details"] = {
                "total_rows": full_shape[0],
                "total_columns": full_shape[1],
                "columns": list(df_sample.columns),
                "dtypes": {col: str(dtype) for col, dtype in df_sample.dtypes.items()},
                "sample_rows": df_sample.fillna("").to_dict(orient="records")
            }

        # 3. معالجة مستندات وورد (Word .docx)
        elif ext == '.docx' and docx:
            meta["type"] = "word"
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            meta["details"] = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "sample_text": paragraphs[:3] if paragraphs else []
            }

        # 4. معالجة ملفات الـ (PDF)
        elif ext == '.pdf' and PyPDF2:
            meta["type"] = "pdf"
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                num_pages = len(reader.pages)
                first_page_text = reader.pages[0].extract_text() if num_pages > 0 else ""
                meta["details"] = {
                    "total_pages": num_pages,
                    "sample_text_page_1": first_page_text[:400] if first_page_text else ""
                }

    except Exception as e:
        meta["error"] = str(e)

    return meta

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        sys.exit(1)

    target_path = sys.argv[1]
    result = extract_metadata(target_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
